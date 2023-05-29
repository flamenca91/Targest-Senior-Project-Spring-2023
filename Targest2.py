# This program searches for red colored tags in the documents as a to begin to extract data
# Instructions on how to use the program:
# 1. run program
# 2. choose the data.txt file containing the path of the documents

# from debug import debug
import logging
import docx
from docx.shared import RGBColor
from tkinter import filedialog
from typing import Tuple
import re
import copy

# Set up the logger for catching errors
logging.basicConfig(level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

doNothing = 0
dicts2Copy = {}     # This will hold the dicts2 content in all documents
parents2Copy = []   # parents2 list copy
filtered_L = []     # Will store the ones without a child tag
filtered_LCopy = []
fullText2Copy = []
parents2 = []       # list of parent tags or child tags

# creates a dict for parent and child tags
dicts = {}
OrphanChild2 = []
orphanChildren2Copy = []
dicts10 = {}
dicts3 = {}         # will hold parentTag and text, Orphan tags
dicts2 = {}         # will hold parentTag and text
orphanDicts = {}    # orphan dictionary
parents9 = []

# declaring different lists that will be used to store, tags and sentences
parentTags = []
parent = []         # This will be used to store everything
child = []          # Used to Store child tags
noChild = []        # Used to Store parentTags with no child
withChild = []      # Used to Store parentTags with child tag
parents = []        # Will be used for future function
orphanTagText = []  # Will be used to hold text of orphanChildTags

# reads the text in the document and use the getcoloredTXT function to get the colored text
def readtxt(filename, color: Tuple[int, int, int]):
    try:
        doc = docx.Document(filename)
        text10 = ""
        fullText = []
        new = []
        global everything
        everything = []  # list of tags and text

        for para in doc.paragraphs:
            # Getting the colored words from the doc
            if (getcoloredTxt(para.runs, color)):
                # Concatenating list of runs between the colored text to single a string
                sentence = "".join(r.text for r in para.runs)
                if len(sentence) > 5:  # this chcekcts if sentence has atleast 5 characters
                    fullText.append(sentence)
                # print(sentence) # Prints everything in the terminal
                everything.append(sentence)
                text10 = sentence
                parent.append("".join(r.text for r in para.runs))

        global hasChild  # Will store the ones with a child tag
        global fullText2  # will store everything found
        global children

        global orphanss
        orphanss = []

        global orphanChildren2  # Will store the orphan child tags for orphanReport
        orphanChildren2 = []
        # Finds the lines without a parentTag
        filtered_L = [value for value in fullText if "[" not in value]

        filtered_L = [s.replace(": ", ":") for s in filtered_L]
        # Finds the lines with a parentTag
        filtered_LCopy.extend(filtered_L)
        hasChild = [value for value in fullText if "[" in value]
        # will store everything found
        fullText2 = [value for value in fullText]
        fullText2 = [s.replace(": ", ":") for s in fullText2]
        fullText2 = [s.replace("[ ", "[") for s in fullText2]
        fullText2 = [s.replace("] ", "]") for s in fullText2]
        fullText2Copy.extend(fullText2)

        return fullText, filtered_L, hasChild, filtered_LCopy, fullText2Copy, fullText2

    except Exception as e:
        # Log an error message
        logging.error('readtxt(): ERROR', exc_info=True)


def getcoloredTxt(runs, color):
    coloredWords, word = [], ""
    try:
        for run in runs:
            if run.font.color.rgb == RGBColor(*color):
                word += str(run.text)  # Saves everything found

            elif word != "":  # This will find the parentTags
                coloredWords.append(word)
                parentTags.append(word)
                parents.append(word)
                word = ""

        if word != "":  # This will find the parentTags
            coloredWords.append(word + "\n")
            # word = removeAfter(word)
            child.append(word)
            withChild.append(word)

    except Exception as e:
        logging.error('getColoredText(): ERROR', e)
    else:
        # Log a success message
        logging.info('getColoredText(): PASS')

    return coloredWords  # returns everything found

def generateReport():  # Will generate the report for tags
    try:
        global filepath
        global filepath2
        global child_tags_list
        global filename_collection      # Created for testing filenames of the open files
        global parents_list  # Created for testing the list of trailing tags in each file
        filename_collection = []
        parents_list = []
        child_tags_list = []
        # create a similar method for opening a folder
        filepath = filedialog.askopenfilename(initialdir="/",
                                              title="",
                                              filetypes=(("text file", "*.txt"),
                                                         ("all files", "*.*")))
        file = open(filepath, 'r')
        file.close()
        # Will store the filepath to the document as a string
        filepath2 = str(filepath)
        a = (filepath2)
        with open(a) as file_in:
            lines = []
            for line in file_in:
                lines.append(line)
        for line2 in lines:
            #print(line2)
            line3 = str(line2)
            line4 = line3.replace('\\', '/')
            line5 = line4.replace('"', '')
            line6 = line5.replace("\n", "")
            # print(line6) plots the path of docx files to open
            fullText = readtxt(filename=line6,
                               color=(255, 0, 0))
            #print(line4)

            fullText10 = str(fullText)
            s = ''.join(fullText10)
            w = (s.replace(']', ']\n\n'))
            filepath3 = str(line4.rsplit('/', 1)[-1])  # change filepath to something.docx
            filepath3 = filepath3.split('.', 1)[0]  # removes .docx of the file name
            filename_collection.append(filepath3)       # Appends all files that are open for testing ******************
            print(filepath3 + " added to the report")
            nameOfDoc = (filepath3 + " added to the report\n")


            # w will be used in the future
            w = (w.replace('([', ''))
            w = (w.replace(',', ''))
            w = (w.replace('' '', ''))
            e = 0

            child2 = removeAfter(child)  # removes everything after the parent tag if there is anything to remove
            # while loop until all the  parentTags has been added to the report

            parents2 = copy.deepcopy(parentTags)  # copy of parent tags list
            parents2Copy.extend(parents2)
            childCopy = copy.deepcopy(child2)
            noParent = []
            noParent2 = []
            global orphanChild
            orphanChild = []
            orphanChildParent = []
            parents9000 = []

            parents2 = [s.replace(" ", "") for s in parents2]  # gets rid of space
            parents_list.append(parents2)  # Appends all the trailing tags for testing   *******************************
            while parentTags:
                noParent.append(parentTags[0])

                for ch in child2:
                    if "[" not in ch:
                        child2.remove(ch)

                if e < len(fullText2):  # as long as variable e is not higher than the lines in fullText2
                    if fullText2[e] in filtered_LCopy:  # filtered_L contains the child tags without a parent tag
                        orphanChild.append(parentTags[0])
                        orphanChildren2.append(parentTags[0])
                        parentTags.remove(parentTags[0])  # Removes that tag after use
                        noParent2.append(" ")
                        parents9000.append(" ")
                        orphanChildParent.append(" ")

                        if child2:
                            if "[" not in child2[0]:  # if it is not a parent tag
                                child2.remove(child2[0])  # Removed that tag from the list
                        e += 1

                    elif fullText2[e] not in filtered_LCopy:
                        parentTags.remove(parentTags[0])  # Removes that tag after use
                        if child2:
                            if "[" not in child2[0]:
                                child2.remove(child2[0])  # Removed that tag from the list

                            parents9000.append(child2[0])
                            noParent.append(child2[0])
                            child2.remove(child2[0])  # Removed that tag from the list
                            e += 1

            parents9.extend(parents9000)
            orphanChildren2Copy.extend(orphanChildren2)

            # Make sure everything is cleared before the program gets the next document
            child2.clear()
            parentTags.clear()
            child.clear()

            global dicts11
            dicts11 = dict(zip(parents2, childCopy))  # creates a dictionary if there is a child tag and parent tag
            dicts.update(dicts)

            noParent = [s.replace(" ", "") for s in noParent]
            orphanChild = [s.replace(" ", "") for s in orphanChild]
            dicts9000 = dict(zip(orphanChild, orphanChildParent))  # orphan dictionary
            orphanDicts.update(dicts9000)
            OrphanChild2.extend(orphanChild)

            text2 = removeParent(everything)  # child tag and text
            text3 = removechild(text2)  # only text list
            text4 = removeText(text2)  # child tags

            parents9000 = [x.strip(' ') for x in parents9000]
            # dicts3 = dict(zip(parents2, childCopy))
            dicts3 = dict(zip(parents2, parents9000))
            dicts10.update(dicts3)
            dicts2 = dict(zip(parents2, text3))  # creates a dictionary with child tags and text
            dicts100 = copy.deepcopy(dicts2)

            sorted(dicts2.keys())  # sorts the keys in the dictionary
            dicts2Copy.update(dicts100)
            child_tags_list.append(parents2)    # Added for testing child_tags *****************************************

        return filepath2, filtered_L, orphanChild
        #return parents2, dicts2, dicts10, dicts2Copy, parents2Copy, fullText2, filtered_LCopy, dicts3, orphanDicts, OrphanChild2

    except Exception as e:
        # Log an error message
        logging.error('generateReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport(): PASS')


def generateReport2():
    try:

        global dicts11111  # Will be used for the excel report later for child - parent
        dicts11111 = {}
        dicts11111 = copy.deepcopy(dicts10)


        # counters for Excel report2
        global counter1
        counter1 = 2
        global counter2
        counter2 = 1
        global counter3
        counter3 = 0
        global cell
        cell = 0;
        global cell2
        cell2 = 0;

        pattern = r'\[([^\]]+)\]'
        for key in dicts10:
            if type(dicts10[key]) == str:
                matches2 = re.findall(pattern, (dicts10[key]))
            if len(matches2) > 1:

                dicts10[key] = []
                parents2 = []
                for match in matches2:
                    parents2.append(match)

                for tag in parents2:
                    tag = (tag.replace(' ', ''))
                    dicts10[key] += [tag]
            else:
                doNothing = + 1

        parents10 = []  # list of all the parent tag tags
        for value11 in dicts10.values():
            # if the value is a list, extend the parents list with the list
            if isinstance(value11, list):
                parents10.extend(value11)

            # if the value is not a list, append the value to the parents list
            else:
                parents10.append(value11)

        # create a list of all the keys in the dictionary (all child tags)
        values_list = list(dicts2Copy.keys())

        #  creates a list of all the child tags that are not in the parents list
        global childless
        childless = []

        # for loop to check if the child tag is in the parents list
        for element in values_list:
            if "".join(element) not in "".join(parents10):
                childless.append(element)

        # sorts the childless list
        childless.sort()
        print("\n")
        print("childless tag: ")
        # for loop to add the childless tags to the report
        for child0 in childless:
            print(child0)
        print("\n")

        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0

        orphanTagText = removechild(filtered_LCopy)

        while m < len(dicts2Copy):
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1
                duplicates = []
                for key, value in dicts2Copy.items():

                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        if str(stringKey2) in dicts10:  # if the key is in the dictionary
                            text = dicts10[str(stringKey2)]

                        if isinstance(text, list):

                            for tag in text:
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()

                                if (str(tag) in duplicates):
                                    doNothing += 1
                                else:
                                    parentTag1 = ('[' + tag + ']')

                                    cell = str('A' + str(counter1))
                                    cell2 = str(str(parentTag1))
                                    counter1 += 2  # counter for excel report
                                    counter2 += 1  # counter for excel report

                                    if "TBV:" in parentTag1:
                                        print("TBV found")
                                    print("parent: ", parentTag1)
                                    tag.strip()
                                    duplicates.append(str(tag))

                                    for x in PTags:
                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                if "TBV:" in parentTag1:
                                                    print("TBV found")

                                                print(dicts2Copy[str(keyCheck4)])
                                        else:
                                            print("Requirement text not found")
                                            orphanChildren2Copy.append(str(keyCheck4))
                                        for b in PTags:
                                            b = (b.replace(']', ''))

                                            if b == tag:
                                                i += 1
                                                hx = tag
                                                keys = [h for h, v in dicts10.items() if
                                                        check_string(hx, v)]  # finds all the child tags
                                                k += 1
                                                for item in keys:  # keys are child tags of hx/the parent tag

                                                    if item != "" and item != " ":
                                                        print("child: ", item)
                                                        print(dicts2Copy[str(item)])

                                                        if "TBV:" in parentTag1:
                                                            print("TBV found")
                                                        counter2 = counter1 - 1
                                                        cell = str('B' + str(counter2))
                                                        cell2 = str(item)
                                                        counter2 += 1
                                                        counter1 += 1
                                                print("\n")
                                                counter2 += 1
                                                counter1 += 1
                        else:
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()
                            hx10 = text
                            hx10 = hx10.replace('[', '')
                            hx10 = hx10.replace(']', '')
                            if (str(hx10) in duplicates):
                                doNothing += 1


                            else:
                                for x in PTags:
                                    keyCheck = (x.replace('[', ''))
                                    keyCheck2 = (keyCheck.replace(']', ''))
                                    keyCheck3 = (keyCheck2.replace(']', ''))
                                    keyCheck4 = (keyCheck3.replace(' ', ''))

                                    print("parent tag: ", x)
                                    if "TBV:" in x:
                                        print("TBV found")

                                    cell = str('A' + str(counter1))
                                    cell2 = str(str(x))
                                    counter1 += 2  # counter for excel report
                                    counter2 += 1  # counter for excel report

                                    if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                        if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                            print(dicts2Copy[str(keyCheck4)])
                                            if "TBD:" in keyCheck4:
                                                print("TBD found")
                                    else:
                                        print("Requirement text not found")
                                        orphanChildren2Copy.append(str(keyCheck4))
                                    for b in PTags:

                                        if b == dicts10[str(stringKey2)]:
                                            i += 1
                                            text.strip()

                                            hx = str(text)
                                            hx = hx.replace('[', '')
                                            hx = hx.replace(']', '')

                                            duplicates.append(str(hx))

                                            keys = [h for h, v in dicts10.items() if check_string(hx, v)]
                                            k += 1
                                            for item in keys:  # keys are child tags of hx/the parent tag

                                                if item != "" and item != " ":
                                                    print("child: ", item)
                                                    print(dicts2Copy[str(item)])
                                                    counter2 = counter1 - 1
                                                    cell = str('B' + str(counter2))
                                                    cell2 = str(item)
                                                    counter2 += 1
                                                    counter1 += 1
                                            print("\n")
                                            counter2 += 1
                                            counter1 += 1

                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1

                        if 1 == 1:
                            for orphantag in orphanChildren2Copy:

                                if orphantag in dicts10:
                                    print("orphantag: ", orphantag)
                                    if orphantag not in duplicates:
                                        kek = 0

                            doNothing += 1
                        if o < len(orphanTagText):
                            doNothing += 1
                        o += 1
                        if i < len(parents2Copy):
                            doNothing += 1
                        i += 1
        orphanGenReport()

    except Exception as e:
        # Log an error message
        logging.error('generateReport2(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport2(): PASS')

def orphanGenReport():
    duplicates = []
    try:
        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0

        orphanTagText = removechild(filtered_LCopy)
        while m < len(dicts2Copy):
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1

                for key, value in dicts2Copy.items():
                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        text = dicts10[str(stringKey2)]

                        if isinstance(text, list):
                            doNothing += 1
                            for tag in text:
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()
                                if (str(tag) in duplicates):
                                    doNothing += 1
                                else:
                                    duplicates.append(str(tag))

                                    for x in PTags:
                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                doNothing += 1
                                        else:
                                            doNothing += 1
                                        for b in PTags:
                                            b = (b.replace(']', ''))

                                            if b == tag:
                                                i += 1
                                                hx = str(tag)
                                                # report3.add_paragraph("I'm here")
                                                keys = [h for h, v in dicts10.items() if
                                                        check_string(hx, v)]  # finds all the child tags
                                                k += 1
                                                for item in keys:  # keys are child tags of hx/the parent tag

                                                    if item != "" and item != " ":
                                                        doNothing += 1
                        else:
                            doNothing = + 1
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()

                            for x in PTags:
                                keyCheck = (x.replace('[', ''))
                                keyCheck2 = (keyCheck.replace(']', ''))
                                keyCheck3 = (keyCheck2.replace(']', ''))
                                keyCheck4 = (keyCheck3.replace(' ', ''))

                                if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                    if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                        doNothing += 1
                                else:
                                    doNothing += 1
                                for b in PTags:

                                    if b == dicts10[str(stringKey2)]:
                                        i += 1
                                        hx = str(dicts10[str(stringKey2)])
                                        keys = [h for h, v in dicts10.items() if
                                                check_string(hx, v)]  # finds all the child tags
                                        k += 1
                                        for item in keys:  # keys are child tags of hx/the parent tag
                                            if item != "" and item != " ":
                                                 doNothing += 1

                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1
                        if i < len(parents2Copy):
                            orphanss.append(parents2Copy[i])  # adds orphan tags to a list
                            doNothing += 1
                        if o < len(orphanTagText):
                            doNothing += 1
                        o += 1
                        if i < len(parents2Copy):
                            doNothing += 1
                        i += 1

        print("Orphan Tags: ")
        for orph5 in orphanChildren2Copy:
            print(orph5)

        return dicts2Copy

    except Exception as e:
        # Log an error message
        logging.error('orphanReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('orphanReport(): PASS')


def removeParent(text):  # removes parent tags or child tags
    try:
        childAfter = []
        for line in text:
            childAfter = [i.rsplit('[', 1)[0] for i in text]  # removes parent tags
            childAfter = [re.sub("[\(\[].*?[\)\]]", "", e) for e in childAfter]  # removes parent tags that are left
            childAfter = [re.sub("[\{\[].*?[\)\}]", "", e) for e in childAfter]  # removes "pass", "fail", etc.
        return childAfter

    except Exception as e:
        # Log an error message
        logging.error('removeParent(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeParent(): PASS')


def removeText(text6):  # this should remove everything before the parent tag
    try:
        childAfter = [s.split(None, 1)[0] if len(s.split(None, 1)) >= 2 else '' for s in text6]
        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeText(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeText(): PASS')


def removeAfter(childtags):  # removes everything after the  tag, example "pass"
    try:
        seperator = ']'
        childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]

        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeAfter(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeAfter(): PASS')


def removechild(text):  # removes child, this one needs fixing
    try:
        mylst = []
        mylst = [s.split(None, 1)[1] if len(s.split(None, 1)) >= 2 else '' for s in text]
        return mylst
    except Exception as e:
        # Log an error message
        logging.error('removechild(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removechild(): PASS')


def check_string(string1, string2):  # checks if a string1 is in string2
    if isinstance(string2, str):
        string2 = [string2]
    pattern = r'{}(?!\d)'.format(re.escape(string1))
    for s in string2:
        match = re.search(pattern, s)
        if match is not None:
            return True
    return False

if __name__ == "__main__":
    info = generateReport()   # Calls the generateReport
    print(info[2])

    generateReport2()  # Calls the generateReport2 function